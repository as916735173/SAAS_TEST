import openpyxl
import requests
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import os  # 新增日期和文件操作依赖

# 配置路径（移除固定REPORT_PATH，改为动态生成）
TEST_CASE_PATH = "test_cases/api/api_test_cases.xlsx"  # 测试用例文件路径
BASE_URL = "https://managerweb.test.ccclubs.com/ccclubs-manager-web/user/login/login"  # 接口基础地址
TOKEN = "5da8df3ec88045ad8082b65ea2b15755"  # 固定token值

def read_test_cases():
    """读取Excel测试用例"""
    wb = openpyxl.load_workbook(TEST_CASE_PATH)
    ws = wb.active
    test_cases = []
    for row in ws.iter_rows(min_row=2, values_only=True):  # 跳过表头（第1行）
        # 处理"接口地址"列中的"同上"，替换为基础地址
        api_url = row[4] if row[4] != "同上" else BASE_URL
        test_cases.append({
            "case_id": row[0],
            "case_name": row[1],
            "method": row[3],
            "url": api_url,
            "params": eval(row[5]),  # 将字符串转为字典（需确保Excel中参数格式正确）
            "expected": row[6]
        })
    return test_cases

def send_request(method, url, params):
    """发送接口请求并返回响应（携带固定token）"""
    try:
        # 新增：构造包含token的请求头
        headers = {
            'token': '5245fd1ac95144608567bfbba9593a68',
            "Content-Type": "application/json"
        }
        if method.upper() == "POST":
            # 修改：添加headers参数
            response = requests.post(url, json=params, headers=headers)
        else:
            return {"status": "error", "message": "不支持的请求方法"}
        response.raise_for_status()  # 检查HTTP错误状态码
        return {"status": "success", "data": response.json()}
    except Exception as e:
        return {"status": "error", "message": str(e)}

def validate_response(actual, expected):
    """验证响应是否符合预期"""
    # 简单验证逻辑：检查预期关键字是否在实际响应中
    # 可根据实际需求扩展（如精确匹配字段值）
    actual_str = str(actual)
    return all(keyword in actual_str for keyword in expected.split(","))

def generate_report(test_results):
    """生成Word测试报告（优化：文件名添加日期和序号）"""
    doc = Document()
    
    # 标题
    title = doc.add_heading("登录接口测试报告", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 测试概述
    doc.add_heading("测试概述", level=1)
    total = len(test_results)
    passed = sum(1 for res in test_results if res["result"] == "通过")
    doc.add_paragraph(f"总用例数：{total}，通过数：{passed}，失败数：{total - passed}")
    
    # 详细结果
    doc.add_heading("测试详情", level=1)
    table = doc.add_table(rows=1, cols=5, style="Table Grid")
    headers = ["用例编号", "用例名称", "执行结果", "预期响应", "实际响应"]
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True  # 表头加粗
    
    for result in test_results:
        row_cells = table.add_row().cells
        row_cells[0].text = result["case_id"]
        row_cells[1].text = result["case_name"]
        row_cells[2].text = result["result"]
        row_cells[3].text = result["expected"]
        row_cells[4].text = str(result["actual"])
    
    # 新增：动态生成带日期和序号的报告名
    report_dir = "test_reports"
    current_date = datetime.now().strftime("%Y%m%d")  # 格式：20240610
    base_name = f"登录接口测试报告_{current_date}_"
    
    # 确保目录存在
    os.makedirs(report_dir, exist_ok=True)
    
    # 查找已存在的报告文件，确定序号
    existing_files = [f for f in os.listdir(report_dir) 
                      if f.startswith(base_name) and f.endswith(".docx")]
    max_num = 0
    for file in existing_files:
        try:
            # 从文件名提取序号（如"登录接口测试报告_20240610_3.docx"提取3）
            num = int(file.split("_")[-1].split(".")[0])
            if num > max_num:
                max_num = num
        except:
            continue  # 跳过格式异常的文件
    
    # 生成新文件名（无现有文件时序号为1）
    report_name = f"{base_name}{max_num + 1 if existing_files else 1}.docx"
    report_path = os.path.join(report_dir, report_name)
    
    # 保存报告
    doc.save(report_path)
    print(f"测试报告已生成：{report_path}")

def main():
    # 1. 读取测试用例
    test_cases = read_test_cases()
    if not test_cases:
        print("未读取到测试用例，请检查Excel文件！")
        return
    
    # 2. 执行测试
    test_results = []
    for case in test_cases:
        print(f"正在执行用例：{case['case_id']} - {case['case_name']}")
        # 发送请求（已自动携带token）
        resp = send_request(case["method"], case["url"], case["params"])
        # 处理响应
        if resp["status"] == "error":
            actual = f"请求失败：{resp['message']}"
            result = "失败"
        else:
            actual = resp["data"]
            result = "通过" if validate_response(actual, case["expected"]) else "失败"
        # 记录结果
        test_results.append({
            "case_id": case["case_id"],
            "case_name": case["case_name"],
            "result": result,
            "expected": case["expected"],
            "actual": actual
        })
        # 终端输出
        print(f"结果：{result} | 预期：{case['expected']} | 实际：{actual}\n")
    
    # 3. 生成报告
    generate_report(test_results)

if __name__ == "__main__":
    main()